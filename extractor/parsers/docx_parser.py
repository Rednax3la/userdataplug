"""DOCX Parser — extracts text paragraphs and table content."""

import logging
from typing import Any

logger = logging.getLogger(__name__)


def parse_docx(file_path: str) -> list[dict[str, Any]]:
    """Parse a DOCX file and return text + table blocks."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed")
        return []

    blocks = []
    try:
        doc = Document(file_path)

        # Extract paragraphs as a single text block
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            blocks.append({
                "type": "text",
                "content": "\n".join(paragraphs),
                "source": "docx_text",
            })

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                blocks.append({
                    "type": "table",
                    "content": rows,
                    "source": "docx_table",
                })

    except Exception as e:
        logger.error(f"DOCX parse error: {e}")

    return blocks
